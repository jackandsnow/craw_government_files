import random
import time

import docx
import requests
import pandas as pd

from openpyxl import load_workbook

# use different agents to crawl data, avoiding IP banned
agent = [
    'Mozilla/5.0 (Windows; U; Windows NT 5.1; en-US; rv:1.8.1.2pre) Gecko/20070215 K-Ninja/2.1.1',
    'Mozilla/4.0 (compatible; MSIE 7.0b; Windows NT 5.2; .NET CLR 1.1.4322; .NET CLR 2.0.50727; InfoPath.2; .NET CLR 3.0.04506.30)',
    'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/21.0.1180.71 Safari/537.1 LBBROWSER',
    'Mozilla/5.0 (Windows; U; Windows NT 5.1; zh-CN; rv:1.9) Gecko/20080705 Firefox/3.0 Kapiko/3.0',
    'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/535.11 (KHTML, like Gecko) Chrome/17.0.963.56 Safari/535.11',
    'Mozilla/4.0 (compatible; MSIE 6.0; Windows NT 5.1; SV1; AcooBrowser; .NET CLR 1.1.4322; .NET CLR 2.0.50727)',
    'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_7_3) AppleWebKit/535.20 (KHTML, like Gecko) Chrome/19.0.1036.7 Safari/535.20',
    'Opera/9.80 (Macintosh; Intel Mac OS X 10.6.8; U; fr) Presto/2.9.168 Version/11.52',
    'Mozilla/5.0 (Windows NT 6.1; Win64; x64; rv:2.0b13pre) Gecko/20110307 Firefox/4.0b13pre',
    'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/536.11 (KHTML, like Gecko) Chrome/20.0.1132.11 TaoBrowser/2.0 Safari/536.11',
    'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/535.11 (KHTML, like Gecko) Chrome/17.0.963.84 Safari/535.11 LBBROWSER',
    'Mozilla/4.0 (compatible; MSIE 6.0; Windows NT 5.1; SV1; QQDownload 732; .NET4.0C; .NET4.0E)',
    'Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 5.1; Trident/4.0; SV1; QQDownload 732; .NET4.0C; .NET4.0E; 360SE)',
    'Mozilla/4.0 (compatible; MSIE 7.0; AOL 9.5; AOLBuild 4337.35; Windows NT 5.1; .NET CLR 1.1.4322; .NET CLR 2.0.50727)',
    'Mozilla/4.0 (compatible; MSIE 6.0; Windows NT 5.1; SV1; QQDownload 732; .NET4.0C; .NET4.0E)',
    'Mozilla/5.0 (Windows NT 5.1) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/21.0.1180.89 Safari/537.1',
    'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.1 (KHTML, like Gecko) Chrome/21.0.1180.89 Safari/537.1',
    'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:16.0) Gecko/20100101 Firefox/16.0',
    'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.11 (KHTML, like Gecko) Chrome/23.0.1271.64 Safari/537.11',
    'Mozilla/5.0 (X11; U; Linux x86_64; zh-CN; rv:1.9.2.10) Gecko/20100922 Ubuntu/10.10 (maverick) Firefox/3.6.10'
]


def get_html_text(url, params=None, proxies=None, total=3):
    """
    get the text of target url
    :param url: target url
    :param params: params dict, like {'param1': 'value1', 'param2': 'value2'}
    :param proxies: proxies dict, like {'http': 'proxy1', 'https': proxy2}, its keys are unchangeable
    :param total: max repeat times if time out
    :return: html text or None
    """
    try:
        headers = {'User-Agent': random.choice(agent)}
        r = requests.get(url, params=params, proxies=proxies, headers=headers, timeout=5)
    except requests.exceptions.ConnectionError:  # connection error
        print('Connection Error:', url)
        return None
    except Exception:
        if total > 0:
            time.sleep(5)  # after 5 seconds continue to craw
            return get_html_text(url, params, proxies, total - 1)
        return None
    # get encodings of the website
    encodings = requests.utils.get_encodings_from_content(r.text)
    if len(encodings) == 0:
        r.encoding = 'gbk'
    else:
        r.encoding = encodings[0]
    return r.text


def get_total_page_num(url, prefix, suffix, delimiter):
    """
    获取文件通知的总页数
    :param url: target url
    :param prefix: 分割的前缀
    :param suffix: 分割的后缀
    :param delimiter: 分隔符
    :return: int page number
        eg. html structure like '<script>createPageHTML(50, 0, "index","htm",849);</script>'
        then prefix can be 'createPageHTML(', suffix can be ');</script>' and delimiter is ','
    """
    html_page = get_html_text(url)
    prev = html_page.split(prefix)[1]  # 截掉前半部分
    suff = prev.split(suffix)[0]  # 截掉后半部分
    page_num = suff.split(delimiter)[0]  # 分割拿到页数
    # print(page_num)
    return int(page_num)


def download_file(file_url, filename, total=3):
    """
    下载文件
    :param file_url: 文件URL
    :param filename: 文件名
    :param total: 最大下载次数（防止失败）
    :return: bool value
    """
    try:
        res = requests.get(file_url)
        if res.status_code == 200:
            fp = open(filename, mode='wb')
            fp.write(res.content)
            fp.close()
        else:
            print('File Not Found:', file_url)
    except Exception:
        if total > 0:
            return download_file(file_url, filename, total - 1)
        return False
    return True


def write_word_file(filename, title, data_list):
    """
    write text data to word file
    :param filename: word filename, like '/path/filename.doc' or '/path/filename.docx'
    :param title: title for word file, or maybe None
    :param data_list: paragraphs of word content
    :return: None
    """
    doc = docx.Document()
    if title:
        doc.add_heading(title, 0)
    for data in data_list:
        doc.add_paragraph(data)
    doc.save(filename)


def write_excel_file(filename, data_list, sheet_name='Sheet1', columns=None):
    """
    write list data to excel file, supporting add new sheet
    :param filename: excel filename, like '/path/filename.xlsx' or '/path/filename.xls'
    :param data_list: list data for saving
    :param sheet_name: excel sheet name, default 'Sheet1'
    :param columns: excel column names
    :return: None
    """
    writer = pd.ExcelWriter(filename)
    frame = pd.DataFrame(data_list, columns=columns)
    book = load_workbook(writer.path)
    writer.book = book
    frame.to_excel(excel_writer=writer, sheet_name=sheet_name, index=None)
    writer.close()


def filename_replace(filename):
    """
    replace some special characters of filename
    :param filename: filename before replacing
    :return: filename after replacing
    """
    replace = filename.replace('/', '-').replace('\\', '-').replace(' ', '') \
        .replace('<', '(').replace('>', ')').replace('"', '-')
    return replace
